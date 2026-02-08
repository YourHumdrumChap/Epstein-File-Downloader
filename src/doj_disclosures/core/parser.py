from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path

import fitz  # PyMuPDF
from bs4 import BeautifulSoup
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class ParsedDocument:
    title: str
    text: str
    ocr_used: bool


class DocumentParser:
    def __init__(
        self,
        *,
        ocr_enabled: bool,
        ocr_engine: str = "tesseract",
        ocr_dpi: int = 200,
        ocr_preprocess: bool = True,
        ocr_median_filter: bool = True,
        ocr_threshold: int | None = None,
    ) -> None:
        self._ocr_enabled = ocr_enabled
        self._ocr_engine = (ocr_engine or "tesseract").strip().lower()
        self._ocr_dpi = int(ocr_dpi)
        self._ocr_preprocess = bool(ocr_preprocess)
        self._ocr_median_filter = bool(ocr_median_filter)
        self._ocr_threshold = ocr_threshold

    def parse(self, path: Path, content_type: str, fallback_title: str = "") -> ParsedDocument:
        suffix = path.suffix.lower()
        if "pdf" in content_type.lower() or suffix == ".pdf":
            return self._parse_pdf(path, fallback_title=fallback_title)
        if suffix == ".docx":
            return self._parse_docx(path, fallback_title=fallback_title)
        if suffix == ".txt":
            return self._parse_txt(path, fallback_title=fallback_title)
        if suffix in {".html", ".htm"} or "html" in content_type.lower():
            return self._parse_html(path, fallback_title=fallback_title)
        return self._parse_txt(path, fallback_title=fallback_title)

    def _parse_txt(self, path: Path, fallback_title: str) -> ParsedDocument:
        data = path.read_text(encoding="utf-8", errors="ignore")
        return ParsedDocument(title=fallback_title or path.name, text=data, ocr_used=False)

    def _parse_html(self, path: Path, fallback_title: str) -> ParsedDocument:
        html = path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html, "lxml")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        title = (soup.title.text.strip() if soup.title and soup.title.text else (fallback_title or path.name))
        text = "\n".join(s.strip() for s in soup.get_text("\n").splitlines() if s.strip())
        return ParsedDocument(title=title, text=text, ocr_used=False)

    def _parse_docx(self, path: Path, fallback_title: str) -> ParsedDocument:
        doc = DocxDocument(str(path))
        parts = [p.text for p in doc.paragraphs if p.text]
        return ParsedDocument(title=fallback_title or path.name, text="\n".join(parts), ocr_used=False)

    def _parse_pdf(self, path: Path, fallback_title: str) -> ParsedDocument:
        doc = fitz.open(str(path))
        try:
            page_texts: list[str] = []
            for idx, page in enumerate(doc, start=1):
                t = page.get_text("text")
                if t and t.strip():
                    page_texts.append(f"\n[PAGE {idx}]\n{t}")
                else:
                    page_texts.append(f"\n[PAGE {idx}]\n")

            text = "\n".join(p for p in page_texts if p.strip())

            if self._ocr_enabled and self._looks_like_scanned_pages(page_texts):
                ocr_text = self._ocr_pdf(doc)
                if ocr_text.strip():
                    return ParsedDocument(title=fallback_title or path.name, text=ocr_text, ocr_used=True)
            return ParsedDocument(title=fallback_title or path.name, text=text, ocr_used=False)
        finally:
            doc.close()

    @staticmethod
    def _looks_like_scanned_pages(page_texts: list[str]) -> bool:
        # Heuristic: treat as scanned if most pages have very little extractable text.
        if not page_texts:
            return True
        meaningful = 0
        for t in page_texts:
            # Ignore the [PAGE N] header when judging density.
            cleaned = "\n".join(line for line in t.splitlines() if not line.startswith("[PAGE "))
            if len(cleaned.strip()) >= 40:
                meaningful += 1
        return meaningful / max(1, len(page_texts)) < 0.35

    def _ocr_pdf(self, doc: fitz.Document) -> str:
        if not self._ocr_enabled or self._ocr_engine == "none":
            return ""

        if self._ocr_engine != "tesseract":
            logger.warning("Unsupported OCR engine %r; skipping OCR", self._ocr_engine)
            return ""

        try:
            import pytesseract
            from PIL import Image
        except Exception:
            logger.warning("OCR enabled but pytesseract/Pillow not available")
            return ""

        # Ensure the Tesseract executable is available.
        try:
            import shutil

            cmd = shutil.which("tesseract")
            if not cmd:
                # Common Windows install locations.
                candidates = [
                    r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe",
                    r"C:\\Program Files (x86)\\Tesseract-OCR\\tesseract.exe",
                ]
                for c in candidates:
                    if Path(c).exists():
                        cmd = c
                        break
            if not cmd:
                logger.warning(
                    "OCR enabled but tesseract executable not found on PATH. "
                    "Install Tesseract OCR (tesseract.exe) and ensure it is on PATH."
                )
                return ""
            pytesseract.pytesseract.tesseract_cmd = cmd
        except Exception:
            # Best-effort; if this fails, pytesseract will likely raise later.
            pass

        texts: list[str] = []
        dpi = max(72, min(600, int(self._ocr_dpi)))
        for page in doc:
            pix = page.get_pixmap(dpi=dpi)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            img2 = self._preprocess_for_ocr(img)
            texts.append(pytesseract.image_to_string(img2))
        return "\n".join(t.strip() for t in texts if t.strip())

    def _preprocess_for_ocr(self, img):
        if not self._ocr_preprocess:
            return img

        try:
            from PIL import ImageFilter, ImageOps
        except Exception:
            return img

        # Normalize: grayscale + contrast.
        out = img.convert("L")
        out = ImageOps.autocontrast(out)

        # Light denoise helps Tesseract.
        if self._ocr_median_filter:
            out = out.filter(ImageFilter.MedianFilter(size=3))

        threshold = self._ocr_threshold
        if threshold is None:
            threshold = self._otsu_threshold(out)
        threshold = int(max(0, min(255, threshold)))

        # Binarize.
        out = out.point(lambda p: 255 if p > threshold else 0)
        return out

    @staticmethod
    def _otsu_threshold(img_l) -> int:
        # Otsu's method on an L-mode image; returns threshold in [0,255].
        hist = img_l.histogram()
        if not hist or len(hist) < 256:
            return 128
        total = sum(hist)
        if total <= 0:
            return 128

        sum_total = sum(i * hist[i] for i in range(256))
        sum_b = 0.0
        w_b = 0.0
        var_max = -1.0
        thresh = 128
        for t in range(256):
            w_b += hist[t]
            if w_b == 0:
                continue
            w_f = total - w_b
            if w_f == 0:
                break
            sum_b += t * hist[t]
            m_b = sum_b / w_b
            m_f = (sum_total - sum_b) / w_f
            var_between = w_b * w_f * (m_b - m_f) ** 2
            if var_between > var_max:
                var_max = var_between
                thresh = t
        return int(thresh)
